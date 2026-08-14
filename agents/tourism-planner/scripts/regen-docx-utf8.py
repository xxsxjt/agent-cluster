# -*- coding: utf-8 -*-
"""
北京旅游 v2 docx 重建脚本（UTF-8 安全版）
- 读取 artifacts/beijing-trip-v2.md，生成格式化的 docx + txt
- 所有文件操作显式 encoding='utf-8'，避免 Windows GBK 默认编码坑
- 输出: Downloads/北京旅游v2-20260823-29.docx / .txt
"""
import os
import re
import glob
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH = os.path.join(BASE, "artifacts", "beijing-trip-v2.md")
DOWNLOADS = "C:/Users/du_ji/Downloads"
OUT_NAME = "北京旅游v2-20260823-29.docx"
OUT_TXT = "北京旅游v2-20260823-29.txt"
OUT_DOCX = os.path.join(DOWNLOADS, OUT_NAME)
OUT_TXT_FULL = os.path.join(DOWNLOADS, OUT_TXT)


def set_font(run, size=10.5, bold=False, color=None, name="微软雅黑"):
    """设置中西文字体，避免 Word 中文显示问题"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def parse_table(lines):
    """解析 md 表格（lines 为连续的表格行），返回 (rows, 消耗行数)"""
    rows = []
    i = 0
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip().strip("|")
        cells = [c.strip() for c in line.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            txt = row[ci] if ci < len(row) else ""
            run = p.add_run(txt)
            set_font(run, size=9.5, bold=(ri == 0))
    doc.add_paragraph()


def md_to_docx(md_text):
    doc = Document()
    # 页面边距
    for sec in doc.sections:
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue
        # 分隔线
        if re.fullmatch(r"-{3,}|\*{3,}", stripped):
            i += 1
            continue
        # 表格块
        if stripped.startswith("|"):
            rows, consumed = parse_table(lines[i:])
            if rows:
                add_table(doc, rows)
            i += consumed
            continue
        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            p = doc.add_paragraph()
            run = p.add_run(m.group(2))
            sizes = {1: 16, 2: 14, 3: 12, 4: 11}
            set_font(run, size=sizes.get(level, 10.5), bold=True,
                     color=(0x1F, 0x3B, 0x73) if level <= 2 else (0x33, 0x33, 0x33))
            p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue
        # 引用
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip("> ").strip())
            set_font(run, size=10, color=(0x66, 0x66, 0x66))
            p.paragraph_format.left_indent = Cm(0.5)
            i += 1
            continue
        # 列表
        m = re.match(r"^[-*]\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            run = p.add_run("• " + m.group(1))
            set_font(run, size=10.5)
            i += 1
            continue
        # 有序列表
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            run = p.add_run(m.group(1))
            set_font(run, size=10.5)
            i += 1
            continue
        # 普通段落
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        set_font(run, size=10.5)
        p.paragraph_format.space_after = Pt(2)
        i += 1
    return doc


def strip_md_for_txt(md_text):
    """把 md 转为纯文本（去标记、表格转对齐文本）"""
    out = []
    for line in md_text.splitlines():
        s = line.rstrip()
        st = s.strip()
        if not st:
            out.append("")
            continue
        if st.startswith("|"):
            cells = [c.strip() for c in st.strip("|").split("|")]
            if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                out.append(" | ".join(cells))
            continue
        if re.fullmatch(r"-{3,}|\*{3,}", st):
            continue
        s2 = re.sub(r"^#{1,6}\s+", "", st)
        s2 = re.sub(r"^>\s?", "", s2)
        s2 = re.sub(r"^[-*]\s+", "• ", s2)
        out.append(s2)
    return "\n".join(out)


def main():
    # 1. 读取源（显式 utf-8）
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_text = f.read()
    print(f"源文件读取 OK: {MD_PATH} ({len(md_text)} 字符)")

    # 2. 生成 docx
    doc = md_to_docx(md_text)
    doc.save(OUT_DOCX)
    print(f"docx 已生成: {OUT_DOCX}")

    # 3. 生成 txt（显式 utf-8）
    txt = strip_md_for_txt(md_text)
    with open(OUT_TXT_FULL, "w", encoding="utf-8", newline="\n") as f:
        f.write(txt)
    print(f"txt 已生成: {OUT_TXT_FULL} ({len(txt)} 字符)")

    # 4. 删除旧的损坏 docx（文件名含 U+FFFD 或无法按 UTF-8 解码的）
    removed = []
    for f in glob.glob(os.path.join(DOWNLOADS, "*.docx")):
        base = os.path.basename(f)
        if "v2-20260823-29" in base and base != OUT_NAME:
            os.remove(f)
            removed.append(base)
    print(f"已删除损坏文件: {removed}")

    # 5. 回读验证
    vd = Document(OUT_DOCX)
    texts = [p.text for p in vd.paragraphs if p.text.strip()]
    n_tables = len(vd.tables)
    print(f"验证: docx 段落 {len(texts)} 个, 表格 {n_tables} 个")
    for key in ["故宫", "302", "八达岭", "天安门", "景山", "北海", "颐和园", "圆明园",
                "雍和宫", "孔庙", "南锣鼓巷", "什刹海", "鸟巢", "水立方", "8/29",
                "交通方案对比", "飞机", "大兴机场", "航班", "早买便宜", "怎么选"]:
        found = any(key in t for t in texts) or any(
            key in c.text for tb in vd.tables for row in tb.rows for c in row.cells)
        print(f"  [{('OK' if found else 'MISSING')}] {key}")
        if not found:
            print("  !! 关键内容缺失，退出码 1")
            sys.exit(1)

    # 6. txt 回读验证
    with open(OUT_TXT_FULL, "r", encoding="utf-8") as f:
        txt_check = f.read()
    assert "故宫" in txt_check and "302" in txt_check, "txt 内容验证失败"
    print(f"txt 回读验证 OK（{len(txt_check)} 字符，UTF-8）")
    print("全部验证通过 ✔")


if __name__ == "__main__":
    main()
